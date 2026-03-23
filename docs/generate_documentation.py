"""
Generate formal documentation as Word (.docx) and PDF in this folder.
Run from project root: python docs/generate_documentation.py
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


HERE = Path(__file__).resolve().parent
DOC_TITLE = "Enterprise Policy Search Agent"
VERSION = "1.0"


def add_docx() -> None:
    doc = Document()
    title = doc.add_heading(DOC_TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Technical documentation")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(11)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version {VERSION}  |  {date.today().isoformat()}")

    doc.add_paragraph()

    def h(text: str, level: int = 1):
        doc.add_heading(text, level=level)

    def p(text: str):
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(6)

    h("1. Introduction", 1)
    p(
        "This document describes the Enterprise Policy Search Agent: a web application "
        "built with FastAPI that searches and answers questions over indexed enterprise content "
        "using Google Cloud Vertex AI Search (Discovery Engine). An optional ingestion pipeline "
        "can pull files from Google Drive into Cloud Storage and use Document AI for OCR when "
        "PDFs are image-based (scanned)."
    )

    h("2. System overview", 1)
    p(
        "The application exposes a browser UI (Jinja2 templates and static assets) and JSON APIs. "
        "Search uses the Discovery Engine Search API; conversational answer generation uses "
        "answer_query with citations mapped back to retrieved documents. A grounding status "
        "(strong, partial, or weak) reflects overlap between the user query and retrieved text."
    )
    bullets = [
        "POST /api/search — ranked document hits with snippets and extractive segments.",
        "POST /api/ask — generated answer, citations, references, and grounding_status.",
        "GET / — interactive web UI.",
        "GET /health — liveness check.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    h("3. Prerequisites", 1)
    pre = [
        "Python 3.10 or newer (recommended).",
        "A Google Cloud project with Vertex AI Search configured and a datastore "
        "(for example documents in Google Cloud Storage).",
        "Relevant APIs enabled: Discovery Engine; optionally Document AI, Cloud Storage, "
        "and Google Drive for ingestion.",
        "Application Default Credentials for API access (for example "
        "gcloud auth application-default login locally, or a service account in production).",
    ]
    for line in pre:
        doc.add_paragraph(line, style="List Bullet")

    h("4. Configuration", 1)
    p(
        "Create a file named .env in the project root. The application loads variables with "
        "python-dotenv. PROJECT_ID and APP_ID are required."
    )
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Variable", "Description"),
        ("PROJECT_ID", "Google Cloud project ID."),
        ("APP_ID", "Vertex AI Search engine ID."),
        ("LOCATION", "Engine location (default: global). Use the region that matches your engine."),
        ("PORT", "HTTP port for the web server (default: 8000)."),
    ]
    for i, (a, b) in enumerate(rows):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b

    p(
        "Serving configuration resource names are built in app/vertex_search.py: "
        "default_search for search and default_serving_config for answers. Adjust these "
        "names if your engine uses different serving configs."
    )

    h("5. Installation and operation", 1)
    p("From the project root, using a virtual environment:")
    for line in [
        "python -m venv .venv",
        r".venv\Scripts\activate   (Windows)",
        "pip install -r requirements.txt",
        "python -m app.main",
    ]:
        doc.add_paragraph(line, style="List Number")

    p(
        "Alternatively run uvicorn: uvicorn app.main:app --reload --host 127.0.0.1 --port 8000 "
        "(or the port set in PORT). Open http://127.0.0.1:8000 in a browser."
    )

    h("6. API summary", 1)
    t2 = doc.add_table(rows=3, cols=3)
    t2.style = "Table Grid"
    api_rows = [
        ("Method and path", "Request body (JSON)", "Response highlights"),
        (
            "POST /api/search",
            "question; optional user_pseudo_id",
            "results: array of document hits",
        ),
        (
            "POST /api/ask",
            "question; optional user_pseudo_id, session_id",
            "answer, citations, references, grounding_status",
        ),
    ]
    for i, row in enumerate(api_rows):
        for j, cell in enumerate(row):
            t2.rows[i].cells[j].text = cell

    h("7. Optional ingestion pipeline", 1)
    p(
        "Module app/ingestion/ingest_drive_to_gcs.py implements OAuth to Google Drive, "
        "downloads or exports files from a configured folder, archives raw objects under "
        "archive/processed/ in a bucket, and writes processed outputs under processed/. "
        "PDFs are classified as scanned or text-based; scanned PDFs are sent to Document AI "
        "and stored as .txt. Constants (folder ID, bucket, processor ID, project) are at "
        "the top of the script and must be set for your environment."
    )
    p("Run: python -m app.ingestion.ingest_drive_to_gcs")

    h("8. Project structure", 1)
    for line in [
        "app/main.py — FastAPI application, routes, static files, templates",
        "app/config.py — environment configuration",
        "app/vertex_search.py — Discovery Engine search and answer generation",
        "app/templates/, app/static/ — web UI",
        "app/ingestion/ — Drive to GCS and OCR helpers",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    h("9. Security", 1)
    p(
        "Do not commit .env, credentials.json, token.json, or service account keys. "
        "Treat user_pseudo_id and session_id as identifiers under your governance for "
        "Vertex AI Search analytics and sessions."
    )

    out = HERE / "Enterprise_Policy_Search_Agent.docx"
    doc.save(out)
    print(f"Wrote {out}")


def _pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="BodyJustify",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            alignment=TA_JUSTIFY,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Heading1Custom",
            parent=styles["Heading1"],
            fontSize=14,
            leading=18,
            spaceBefore=12,
            spaceAfter=8,
        )
    )
    return styles


def add_pdf() -> None:
    styles = _pdf_styles()
    path = HERE / "Enterprise_Policy_Search_Agent.pdf"
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    story = []

    story.append(Paragraph(f"<b>{DOC_TITLE}</b>", styles["Title"]))
    story.append(Spacer(1, 0.3 * cm))
    story.append(
        Paragraph(
            f"<para align=center>Technical documentation<br/>Version {VERSION} &mdash; {date.today().isoformat()}</para>",
            styles["Normal"],
        )
    )
    story.append(Spacer(1, 0.8 * cm))

    sections = [
        (
            "1. Introduction",
            "This document describes the Enterprise Policy Search Agent: a web application built with "
            "FastAPI that searches and answers questions over indexed enterprise content using Google Cloud "
            "Vertex AI Search (Discovery Engine). An optional ingestion pipeline can pull files from Google "
            "Drive into Cloud Storage and use Document AI for OCR when PDFs are image-based (scanned).",
        ),
        (
            "2. System overview",
            "The application exposes a browser UI and JSON APIs. Search uses the Discovery Engine Search API; "
            "answer generation uses answer_query with citations mapped to retrieved documents. A grounding status "
            "(strong, partial, or weak) reflects query-to-evidence overlap. Endpoints: POST /api/search, POST /api/ask, "
            "GET / (UI), GET /health.",
        ),
        (
            "3. Prerequisites",
            "Python 3.10+ recommended; a GCP project with Vertex AI Search and a datastore; required APIs enabled; "
            "Application Default Credentials for local or production use.",
        ),
        (
            "4. Configuration",
            "Create .env in the project root with PROJECT_ID and APP_ID (required), LOCATION (default global), "
            "and PORT (default 8000). Adjust serving config names in app/vertex_search.py if your engine differs "
            "from default_search and default_serving_config.",
        ),
        (
            "5. Installation and operation",
            "Create a venv, pip install -r requirements.txt, then run python -m app.main or uvicorn app.main:app. "
            "Browse to http://127.0.0.1:8000 (or configured PORT).",
        ),
        (
            "6. API summary",
            "POST /api/search accepts question and optional user_pseudo_id; returns results. "
            "POST /api/ask accepts question and optional user_pseudo_id and session_id; returns answer, citations, "
            "references, and grounding_status.",
        ),
        (
            "7. Optional ingestion",
            "app/ingestion/ingest_drive_to_gcs.py: Drive OAuth, export/download, GCS upload, PDF scan detection, "
            "Document AI OCR for scanned PDFs. Configure constants in the script. Run: python -m app.ingestion.ingest_drive_to_gcs.",
        ),
        (
            "8. Project structure",
            "main.py (app), config.py, vertex_search.py, templates/, static/, ingestion/.",
        ),
        (
            "9. Security",
            "Do not commit secrets. Protect user_pseudo_id and session_id per your policies.",
        ),
    ]

    for title, body in sections:
        story.append(Paragraph(f"<b>{title}</b>", styles["Heading1Custom"]))
        story.append(Paragraph(body, styles["BodyJustify"]))

    data = [
        ["Variable", "Description"],
        ["PROJECT_ID", "Google Cloud project ID"],
        ["APP_ID", "Vertex AI Search engine ID"],
        ["LOCATION", "Engine location (e.g. global)"],
        ["PORT", "HTTP port (default 8000)"],
    ]
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("<b>Configuration variables (.env)</b>", styles["Heading1Custom"]))
    t = Table(data, colWidths=[4 * cm, 12 * cm])
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(t)

    doc.build(story)
    print(f"Wrote {path}")


def main() -> None:
    add_docx()
    add_pdf()


if __name__ == "__main__":
    main()
